from pyautogui import hotkey
from time import sleep
import os
import subprocess
import shutil
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt, RGBColor
from docx2pdf import convert

palavras_minusc = ['de', 'do', 'dos', 'das', 'da']
continuar = 'S'
continuarr = 'S'
primeira_execucao = True

while continuarr != 'N':
    continuar = 'S'
    primeira_execucao = True
    
    while continuar != 'N':
        if primeira_execucao:
            def l():
                print('\033[32m-=-\033[m' * 20)

            def substituir_palavras(paragrafo, palavras_substituir):
                for busca, substituicao in palavras_substituir.items():
                    if busca in paragrafo.text:
                        for run in paragrafo.runs:
                            run.text = run.text.replace(busca, substituicao)

            def obter_sexo():
                while True:
                    l()
                    sexo = input('Digite o sexo do personagem principal (Masculino/Feminino): ').upper().strip()
                    l()
                    if sexo in ['M', 'F']:
                        return sexo
                    else:
                        print('Opção inválida. Por favor, informe M para masculino ou F para feminino.')

            l = lambda: print('-'*100)
            l()
            print(' '*30, 'MEU LIVRO PERSONALIZADO')
            l()
            nome = input('Digite o nome do personagem principal: ').title()
            sobrenome = input('Digite o sobrenome do personagem principal: ')
            amg1 = input('Digite o nome do amigo 1: ').title()
            amg2 = input('Digite o nome do amigo 2: ').title()
            amg3 = input('Digite o nome do amigo 3: ').title()
            cidade = input('Digite a cidade: ').title()
            dedicatoria = input('Digite a dedicatória do livro: ').title()
            qmofrc = input('Digite quem oferece o livro: ').title()
            sexo = obter_sexo()
            pasta = 'Livros Meninos' if sexo == 'M' else 'Livros Meninas'
            primeira_execucao = False

        if continuar == 'S':
            # Lista todos os arquivos na pasta selecionada
            files = os.listdir(pasta)

            # Filtra apenas os arquivos com extensão .docx
            docx_files = list(filter(lambda x: x.endswith('.docx'), files))

            if len(docx_files) == 0:
                print('Nenhum arquivo .docx encontrado na pasta selecionada.')
            else:
                print('Escolha o livro desejado:')
                l()
                for i, file in enumerate(docx_files):
                    print(f'{i+1}. {file}')
                l()
                while True:
                    escolha = input('Informe o Nº correspondente ao livro desejado: ')
                    if escolha.isdigit() and int(escolha) <= len(docx_files):
                        break
                    else:
                        print('Opção inválida. Por favor, escolha um número válido.')

            # Obtém o arquivo escolhido com base na escolha do usuário
            chosen_file = docx_files[int(escolha) - 1]
            print(f'Arquivo escolhido: {chosen_file}')

            doc_path = os.path.join(pasta, chosen_file)
            if not os.path.isfile(doc_path):
                print(f'O arquivo {chosen_file} não foi encontrado na pasta selecionada.')
                continue

            try:
                doc = Document(doc_path)
            except FileNotFoundError as e:
                print(f'Erro ao abrir o arquivo {chosen_file}: {str(e)}')
                continue

            # Abre o arquivo .docx selecionado
            doc = Document(os.path.join(pasta, chosen_file))

            # Copia o arquivo .docx para um novo arquivo chamado "livro_copia.docx"
            doc_copy_path = os.path.join(pasta, 'livro_copia.docx')
            shutil.copy2(os.path.join(pasta, chosen_file), doc_copy_path)

            # Abre a cópia do arquivo .docx
            doc_copy = Document(doc_copy_path)

            # Renomeia o arquivo com o novo nome
            novo_nome = f'{nome}_{sobrenome}.docx'
            novo_nome_path = os.path.join(pasta, novo_nome)
            if os.path.exists(novo_nome_path):
                os.remove(novo_nome_path)
            os.rename(doc_copy_path, novo_nome_path)

            # Função para substituir palavras no documento
            def substituir_palavra(paragrafo, palavra_antiga, palavra_nova):
                if palavra_antiga in paragrafo.text:
                    inline = paragrafo.runs
                    for i in range(len(inline)):
                        if palavra_antiga in inline[i].text:
                            texto = inline[i].text.replace(palavra_antiga, palavra_nova)
                            inline[i].text = texto

            # Substitui as palavras no documento
            for paragrafo in doc_copy.paragraphs:
                substituir_palavra(paragrafo, 'PERSONAGEM', nome)
                substituir_palavra(paragrafo, 'Personagem', nome)
                substituir_palavra(paragrafo, '<<nome>>', nome)
                substituir_palavra(paragrafo, 'SOBRENOME', sobrenome)
                substituir_palavra(paragrafo, 'DEDICATORIA', dedicatoria)
                substituir_palavra(paragrafo, 'QUEM OFERECE', qmofrc)
                substituir_palavra(paragrafo, 'CIDADE', cidade)
                substituir_palavra(paragrafo, 'AMIGO 01', amg1)
                substituir_palavra(paragrafo, 'AMIGO 02', amg2)
                substituir_palavra(paragrafo, 'AMIGO 03', amg3)
                substituir_palavra(paragrafo, 'sexo', sexo)

            # Define a cor da fonte para preto
            for paragrafo in doc_copy.paragraphs:
                for run in paragrafo.runs:
                    font_color = run.font.color
                    if font_color is not None and font_color.rgb != RGBColor(0, 0, 0):
                        run.font.color.rgb = RGBColor(0, 0, 0)

            # Salva o arquivo .docx com as palavras substituídas
            doc_copy.save(novo_nome_path)

            # Converte o arquivo .docx para .pdf
            pdf_path = os.path.join('livros prontos', f'{nome}_{sobrenome}.pdf')
            convert(novo_nome_path, pdf_path)

            # Remove o arquivo .docx antigo
            os.remove(novo_nome_path)

            # Abre o arquivo .pdf gerado
            subprocess.Popen([pdf_path], shell=True)

        l()
        continuar = input('Deseja personalizar outro livro para a mesma criança? (S/N): ').upper().strip()
        l()
    continuarr = input('Deseja continuar personalizando livros para outras crianças? (S/N): ').upper().strip()
    l()
    
print('Foi um prazer ajudá-lo a transformar suas histórias favoritas em versões personalizadas.\n Agradecemos por escolher o ''Meu Livro Personalizado'' desejamos a você momentos mágicos em suas leituras futuras! \n F-Teck 2023.')
sleep(5)